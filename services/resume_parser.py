import os
import logging
from typing import Optional
import fitz  # PyMuPDF
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParser:
    """Service for parsing resume files."""
    
    @staticmethod
    def parse_pdf_pymupdf(file_path: str) -> str:
        """
        Parse PDF using PyMuPDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            text = ""
            doc = fitz.open(file_path)
            
            for page in doc:
                text += page.get_text()
            
            doc.close()
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing PDF with PyMuPDF: {str(e)}")
            raise
    
    @staticmethod
    def parse_pdf_pdfplumber(file_path: str) -> str:
        """
        Parse PDF using pdfplumber (fallback method).
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            text = ""
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing PDF with pdfplumber: {str(e)}")
            raise
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Parse DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise
    
    @staticmethod
    def parse_resume(file_path: str) -> str:
        """
        Parse resume file (PDF or DOCX).
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Extracted text
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                # Try PyMuPDF first
                try:
                    text = ResumeParser.parse_pdf_pymupdf(file_path)
                    if text:
                        logger.info(f"Successfully parsed PDF with PyMuPDF: {file_path}")
                        return text
                except Exception as e:
                    logger.warning(f"PyMuPDF failed, trying pdfplumber: {str(e)}")
                
                # Fallback to pdfplumber
                text = ResumeParser.parse_pdf_pdfplumber(file_path)
                logger.info(f"Successfully parsed PDF with pdfplumber: {file_path}")
                return text
                
            elif file_extension == '.docx':
                text = ResumeParser.parse_docx(file_path)
                logger.info(f"Successfully parsed DOCX: {file_path}")
                return text
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def validate_resume_text(text: str) -> bool:
        """
        Validate extracted resume text.
        
        Args:
            text: Extracted text
            
        Returns:
            True if valid, False otherwise
        """
        if not text or len(text.strip()) < 100:
            logger.warning("Resume text is too short or empty")
            return False
        
        return True


# Global resume parser instance
resume_parser = ResumeParser()
